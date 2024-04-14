import pandas as pd
import matplotlib.pyplot as plt
import scipy.stats as st
from docx import Document

# File path of your CSV file
#file_path = "./surgical_factors/clean_data.csv"
file_path = "clean_data.csv"

# Use pandas to read the CSV file and load it into a DataFrame
df = pd.read_csv(file_path)


### FUNCTIONS

# Function to calculate the mean and median
def calculate_stats(df, column_name):
    column_data = df[column_name]
    mean = column_data.mean()
    median = column_data.median()
    std = column_data.std()
    print(f'The mean of {column_name} is {mean}, the median is {median}, and the standard deviation is {std}')
    return mean, median, std

# T-test without a seperator 
def independent_ttest(group1, group2):
    t_stat, p_value = st.ttest_ind(group1, group2)
    return t_stat, p_value

# T-test with a seperator 
def independent_ttest_seperator(df, seperator, column):
    group1 = df[df[seperator] == 0][column]
    group2 = df[df[seperator] == 1][column]
    
    # Drop rows with NaN values from each group
    group1 = group1.dropna()
    group2 = group2.dropna()

    # Perform independent t-test
    t_stat, p_value = st.ttest_ind(group1, group2)

    print(f'tstat = {t_stat}, p_value = {p_value}\n')
    return t_stat, p_value


# Create a binary columns based on thresholds
def create_binary_column(df, source_column, threshold, new_column_name):
    df[new_column_name] = (df[source_column] > threshold).astype(int)
    return df


# Normals frequency plot
def single_column(df, column_name, x_labels=None, plot_show = False):
    # Get all unique values in the column
    unique_values = df[column_name].unique()
    unique_values.sort()

    # Count the frequency of each value in the column
    value_counts = df[column_name].value_counts().reindex(unique_values, fill_value=0)

    # Print statements
    print(value_counts)

    # Remove values with count of 0
    value_counts = value_counts[value_counts != 0]

    if len(value_counts) == 0:
        print("All counts are zero. No bars to plot.")
        return

    # Create the bar plot
    plt.bar(range(len(value_counts)), value_counts.values)

    # Add text on top of each bar
    for i, v in enumerate(value_counts.values):
        plt.text(i, v, str(v), ha='center', va='bottom')

    # Add x-label and y-label
    plt.xlabel(column_name)
    plt.ylabel('Frequency')

    # Set custom x-axis labels if provided
    if x_labels is not None:
        plt.xticks(range(len(value_counts)), x_labels)
    else:
        plt.xticks(range(len(value_counts)))  # Remove tick labels

    # Adjust y-axis limits for better visibility of text on top of bars
    plt.ylim(0, max(value_counts.values) * 1.1)

    # Show the plot
    if plot_show == True:
        plt.show()

    calculate_stats(df, column_name)



# specifically for multiple columns histograms
def multi_columns(df, columns, custom_labels=None, plot_show = False):
    # Extract the specified ethnicity columns from the DataFrame
    ethnicity_data = df[columns]

    # Calculate the sum of each ethnicity column
    counts = ethnicity_data.sum()

    print(columns)
    print(counts)

    # Plot the frequencies
    plt.figure(figsize=(10, 6))  # Adjust the figure size as needed
    plt.bar(range(len(counts)), counts.values)

    # Add text on top of each bar
    for i, v in enumerate(counts.values):
        plt.text(i, v, str(v), ha='center', va='bottom')

    # Add x-label and y-label
    plt.ylabel('Frequency')

    # Set x-axis ticks and labels
    if custom_labels is not None:
        plt.xticks(range(len(counts)), custom_labels, rotation=45, ha='right')
    else:
        plt.xticks(range(len(counts)), counts.index, rotation=45, ha='right')

    # Adjust bottom margin to prevent cutting off bottom labels
    plt.subplots_adjust(bottom=0.15)

    # Show the plot
    plt.tight_layout()  # Adjust the layout to prevent label overlap
    # Show the plot
    if plot_show == True:
        plt.show()



# Double bar plot comparison with a seperator (for data like, JR vs SR workload)
def comparison_binary(df, column_name, x_labels=None, separate_column=None, label_0=None, label_1=None, plot_show = False):
    if separate_column is None:
        print("Separate column not provided. Please provide a column name containing values of 0 or 1.")
        return
    
    # Separate the DataFrame into two groups based on the separate_column
    group_0 = df[df[separate_column] == 0]
    group_1 = df[df[separate_column] == 1]

    # Get all unique values in the specified column
    unique_values = df[column_name].unique()
    unique_values.sort()

    # Count the frequency of each value in the specified column for each group
    value_counts_0 = group_0[column_name].value_counts().reindex(unique_values, fill_value=0)
    value_counts_1 = group_1[column_name].value_counts().reindex(unique_values, fill_value=0)

    # Print statements
    print("Group 0:")
    print(value_counts_0)
    print("\nGroup 1:")
    print(value_counts_1)

    # Remove values with count of 0
    value_counts_0 = value_counts_0[value_counts_0 != 0]
    value_counts_1 = value_counts_1[value_counts_1 != 0]

    if len(value_counts_0) == 0 or len(value_counts_1) == 0:
        print("One or both groups have all counts as zero. No bars to plot.")
        return

    # Set the width of each bar
    bar_width = 0.35

    # Create the bar plot for group 0
    plt.bar(range(len(value_counts_0)), value_counts_0.values, width=bar_width, color='b', alpha=0.5, label=label_0)

    # Create the bar plot for group 1, shifted for side-by-side bars
    plt.bar([x + bar_width for x in range(len(value_counts_1))], value_counts_1.values, width=bar_width, color='r', alpha=0.5, label=label_1)

    # Add text on top of each bar for group 0
    for i, v in enumerate(value_counts_0.values):
        plt.text(i, v, str(v), ha='center', va='bottom')

    # Add text on top of each bar for group 1
    for i, v in enumerate(value_counts_1.values):
        plt.text(i + bar_width, v, str(v), ha='center', va='bottom')

    # Add x-label and y-label
    plt.xlabel(column_name)
    plt.ylabel('Frequency')

    # Set custom x-axis labels if provided
    if x_labels is not None:
        plt.xticks([x + bar_width / 2 for x in range(len(value_counts_0))], x_labels)
    else:
        plt.xticks([])  # Remove tick labels

    # Adjust y-axis limits for better visibility of text on top of bars
    max_y = max(max(value_counts_0.values), max(value_counts_1.values))
    plt.ylim(0, max_y * 1.1)

    # Add legend
    plt.legend()

    # Show the plot
    if plot_show == True:
        plt.show()

    # Calculate statistics for both groups
    stats_group_0 = calculate_stats(group_0, column_name)
    stats_group_1 = calculate_stats(group_1, column_name)

    independent_ttest_seperator(df, separate_column, column_name)


def create_word_document_from_df(df, additional_info_header, additional_info_text, header_columns, text_columns, output_filename):
    # Create a new Document
    doc = Document()
    
    # Iterate through each row of the DataFrame
    for index, row in df.iterrows():
        # Initialize a list to store the parts of the header
        header_parts = []
        
        # Iterate through additional_info_header and header_columns simultaneously
        for info, header_col in zip(additional_info_header, header_columns):
            header_parts.append(f"{info} {row[header_col]}")
        
        # Combine header parts into a single header string
        header = ' - '.join(header_parts)
        
        # Add header as a paragraph to the document
        doc.add_heading(header, level=1)
        
        # Iterate through text columns
        for info, col in zip(additional_info_text, text_columns):
            # Add additional info followed by text content as a paragraph
            text_content = f"{info} - {row[col]}\n"
            doc.add_paragraph(text_content)
        
        # Add two empty paragraphs after all text columns
        doc.add_paragraph('')
    
    # Save the document
    doc.save(output_filename)


